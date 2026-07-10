#!/usr/bin/env python3
"""Build a DOCX technical report for ParaView workflows and case visuals."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "analysis_outputs"
SUITE_DIR = OUT_DIR / "savonius_mrf_paraview_suite"


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def add_image(doc: Document, path: Path, caption: str, width_cm: float = 16.5):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Savonius MRF ParaView Teknik Raporu")
    r.bold = True
    r.font.size = Pt(18)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run(f"Tarih: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    add_heading(doc, "1. Amaç", level=1)
    doc.add_paragraph(
        "Bu raporun amacı, Savonius OpenFOAM case'leri için ParaView tarafında "
        "tekrarlanabilir, teknik olarak savunulabilir ve raporlanabilir bir "
        "görselleştirme standardı oluşturmaktır. Hedef yalnız estetik ekran "
        "görüntüsü almak değil; geometri, akış yönü, wake yapısı, hız açığı ve "
        "yerel mesh davranışını mühendislik kararına destek olacak şekilde sunmaktır."
    )

    add_heading(doc, "2. Kaynak Temeli", level=1)
    doc.add_paragraph("Bu rapor iki bilgi eksenini birleştirir: resmi ParaView dokümantasyonu ve kullanıcı/pratik deneyimleri.")
    add_bullet(doc, "Resmi ParaView CFD eğitimi, slice, stream tracer ve tube tabanlı akış görselleştirmesini temel araçlar olarak önerir.")
    add_bullet(doc, "ParaView topluluk önerilerinde 3D CFD veriyi önce uygun bir slice ile 2D yoruma indirmek, ardından stream tracer uygulamak sık tekrar edilen iyi pratiktir.")
    add_bullet(doc, "OpenFOAM için boş bir .foam dosyasının reader giriş noktası olarak kullanılması topluluk içinde kabul gören standart yöntemdir.")
    add_bullet(doc, "Gradient/vorticity türev alanları, unstructured ve yerel bozuk hücrelerde pahalı veya gürültülü olabilir; bu yüzden derived alanlar körlemesine kullanılmamalıdır.")
    add_bullet(doc, "Streamline okunabilirliği için tube kullanımı resmi dokümanda ve kullanıcı paylaşımlarında tekrar tekrar önerilmektedir.")

    doc.add_paragraph("Kaynak bağlantıları:")
    doc.add_paragraph("https://docs.paraview.org/en/latest/Tutorials/ClassroomTutorials/targetedComputationFluidDynamics.html")
    doc.add_paragraph("https://docs.paraview.org/en/latest/UsersGuide/filteringData.html")
    doc.add_paragraph("https://docs.paraview.org/en/latest/UsersGuide/displayingData.html")
    doc.add_paragraph("https://discourse.paraview.org/t/best-practice-for-plotting-streamlines-around-a-cylinder/2351")
    doc.add_paragraph("https://discourse.paraview.org/t/how-to-compute-vorticity-in-paraview/4474")
    doc.add_paragraph("https://discourse.paraview.org/t/2d-streamline-with-vti-look-strange/9891")
    doc.add_paragraph("https://discourse.paraview.org/t/creating-streamlines/3290")
    doc.add_paragraph("https://discourse.paraview.org/t/increase-line-thickness/13964")

    add_heading(doc, "3. Bu Case İçin Görselleştirme Stratejisi", level=1)
    add_bullet(doc, "Bu Savonius case tam 2D değildir; ince bir 3D symmetry modeldir. Bu yüzden en güvenilir okuma orta düzlem slice üzerinden yapılır.")
    add_bullet(doc, "Akış yönünü göstermek için line-seeded stream tracer kullanıldı. Streamline'lar tube ile kalınlaştırıldı; ince çizgiler raporda zayıf kalıyordu.")
    add_bullet(doc, "Wake yorumu için vorticity yerine velocity deficit kullanıldı. Sebep: mevcut meshte Gradient filtresi türev hesabında Jacobian uyarıları üretiyor ve otomatik rapor pipeline'ını yavaşlatıyor.")
    add_bullet(doc, "3D görünümde amaç hacim doldurmak değil, rotor etrafındaki akış sapmasını ve wake'in uzanımını göstermekti.")
    add_bullet(doc, "Mesh görseli, güzellik için değil; refinement bölgesinin gerçekten rotor çevresinde yoğunlaşıp yoğunlaşmadığını doğrulamak için alındı.")

    add_heading(doc, "4. Üretilen Görseller ve Teknik Yorumu", level=1)
    add_image(doc, SUITE_DIR / "02_midplane_velocity_streamlines.png", "Sekil 1. Orta duzlem hiz buyuklugu ve streamline yapisi. Rotor oncesi hizlanma, rotor arkasi sapma ve wake ayrilmasi goruluyor.")
    add_image(doc, SUITE_DIR / "03_midplane_wake_deficit.png", "Sekil 2. Orta duzlem hiz acigi (6-Ux). Rotor arkasindaki dusuk momentum cekirdegi ve wake genislemesi acik seciliyor.")
    add_image(doc, SUITE_DIR / "04_3d_wake_structure.png", "Sekil 3. 3D streamtube gorunumu. Amaç, rotor etrafinda akisin nasil yon degistirdigini ve wake'e nasil baglandigini gostermek.")
    add_image(doc, SUITE_DIR / "05_local_mesh_near_blade.png", "Sekil 4. Rotor cevresindeki yerel mesh yogunlasmasi. Refine bolgesinin geometri ve wake cekirdegi etrafinda toplandigi dogrulaniyor.")

    add_heading(doc, "5. Basarisiz veya Reddedilen Yaklasimlar", level=1)
    add_bullet(doc, "Derived vorticity haritasi ilk denemede cikarildi ancak domain sinirlarinda sahte parlama ve Gradient filtresi kaynakli pahali hesap maliyeti gozlendi.")
    add_bullet(doc, "3D contour tabanli wake izoyuzeyleri bu meshte gereksiz derecede maliyetliydi; rapor kalitesine katkisi, streamtube + wake-deficit kombinasyonuna gore daha dusuktu.")
    add_bullet(doc, "Blade pressure close-up, mevcut quasi-2D ince geometri nedeniyle ParaView'da tek basina en guclu anlatim araci olmadi; bu alan icin Matplotlib tabanli yuzey haritasi daha okunakli kaldi.")

    add_heading(doc, "6. Case-ozel Teknik Bulgular", level=1)
    add_bullet(doc, "Wake, rotorun hemen arkasinda tek parca bir dusuk momentum bulutu olarak baslayip downstream yonunde genisliyor.")
    add_bullet(doc, "Streamline yapisi rotor ustunde hizlanma, rotor arkasinda ise belirgin sapma ve acisal asimetri gosteriyor.")
    add_bullet(doc, "Mesh, rotor etrafinda yeterli yerel yogunlasmaya sahip; fakat concave-cell sinyali halen oldugu icin derived alanlarda ekstra dikkat gerekiyor.")
    add_bullet(doc, "Bu nedenle bu case icin karar verici gorsellerin sirasi: velocity slice + streamlines > wake deficit > local mesh > derived vortical quantities.")

    add_heading(doc, "7. Operasyonel Standart", level=1)
    add_bullet(doc, "Akis topolojisi icin once orta duzlem slice al.")
    add_bullet(doc, "Ayni sahnede streamline kullanacaksan tube ile kalinlastir.")
    add_bullet(doc, "Wake yorumu icin ilk tercih Ux veya velocity deficit olsun; vorticity ikinci seviye arac olarak kalsin.")
    add_bullet(doc, "3D goruntu alirken tum domain'i doldurma; rotor ve yakin wake bolgesine odaklan.")
    add_bullet(doc, "Mesh gorselini yalniz refinement ve hucre boyutu mantigini kanitlamak icin kullan.")

    add_heading(doc, "8. Dosyalar", level=1)
    doc.add_paragraph("Uretilen ana dosyalar:")
    for path in [
        SUITE_DIR / "02_midplane_velocity_streamlines.png",
        SUITE_DIR / "03_midplane_wake_deficit.png",
        SUITE_DIR / "04_3d_wake_structure.png",
        SUITE_DIR / "05_local_mesh_near_blade.png",
        ROOT / "analysis_scripts" / "paraview_case_suite.py",
    ]:
        doc.add_paragraph(str(path))

    return doc


def main():
    doc = build_doc()
    out = ROOT / "PARAVIEW_TEKNIK_RAPOR_2026-07-10.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
